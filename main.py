# -*- coding: utf-8 -*-

import json
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor


# Syntactic sugar to put formatting options on a run object
def format_run(run, formatting_options):
    for formatting_option in formatting_options:
        setattr(run.font, formatting_option, True)


# Encoding used to read the JSON content file
content_encoding = 'utf-8'

# JSON keys used in settings.json file, in case they would change, just modify those variables
document_settings = 'document_settings'
header_settings = 'header_settings'
word_settings = 'word_settings'
definition_settings = 'definition_settings'
both_word_and_definition_settings = 'both_word_and_definition_settings'

# Basic look-up table for human-readable alignment values
alignments = {
    'LEFT': WD_PARAGRAPH_ALIGNMENT.LEFT,
    'RIGHT': WD_PARAGRAPH_ALIGNMENT.RIGHT,
    'CENTER': WD_PARAGRAPH_ALIGNMENT.CENTER,
    'JUSTIFY': WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
    "DISTRIBUTE": WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
}

with open('settings.json') as config_file, open('content.json', encoding=content_encoding) as content_file:
    configuration = json.load(config_file)
    content = json.load(content_file)['content']

# Document creation from scratch
document = Document()
encoding = configuration[document_settings]['encoding']

# Header part
header_paragraph = document.add_paragraph()
header_run = header_paragraph.add_run(configuration[header_settings]['text'])
header_run.font.name = configuration[header_settings]['font']
header_run.font.size = Pt(configuration[header_settings]['size'])
header_run.font.color.rgb = RGBColor(*tuple(configuration[header_settings]['color']))
header_run.add_break()
format_run(header_run, configuration[header_settings]['formatting'])
header_paragraph.paragraph_format.alignment = alignments[configuration[header_settings]['alignment']]

# List of definitions part
content.sort(key=lambda k: k['word'])
for row in content:
    word, definition = row['word'], row['definition']
    word += ' {} '.format(configuration[both_word_and_definition_settings]['separator'])

    row_paragraph = document.add_paragraph()

    word_run = row_paragraph.add_run(word)
    word_run.font.name = configuration[word_settings]['font']
    word_run.font.size = Pt(configuration[word_settings]['size'])
    word_run.font.color.rgb = RGBColor(*tuple(configuration[word_settings]['color']))
    format_run(word_run, configuration[word_settings]['formatting'])

    definition_run = row_paragraph.add_run(definition)
    definition_run.font.name = configuration[definition_settings]['font']
    definition_run.font.size = Pt(configuration[definition_settings]['size'])
    definition_run.font.color.rgb = RGBColor(*tuple(configuration[definition_settings]['color']))

    row_paragraph.paragraph_format.alignment = alignments[configuration[both_word_and_definition_settings]['alignment']]

# Document saving
document.save(configuration[document_settings]['filename'])
